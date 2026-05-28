from pathlib import Path

from docx import Document


path = Path("论文材料") / "曹逸天-毕业设计说明书-学校模板Word重排版-实验增强版-图表补充版.docx"
doc = Document(path)
print("paragraphs", len(doc.paragraphs), "tables", len(doc.tables), "images", len(doc.inline_shapes))
needles = ["6.8 补充数据统计与结论图表", "图 6-2", "YOLO 外部桥接实验中", "第 7 章"]
for needle in needles:
    hits = [(i, p.text[:120]) for i, p in enumerate(doc.paragraphs) if needle in p.text]
    print("==", needle, "==")
    for hit in hits[:5]:
        print(hit)
