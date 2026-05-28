from pathlib import Path
import os
import re

from docx import Document


DOCX = Path(os.environ.get("FINAL_DOCX", "毕业设计提交归档-曹逸天/3毕业设计说明书-曹逸天-学校模板图表补充精修版.docx"))


def main():
    doc = Document(DOCX)
    headings = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if re.match(r"^(第\s*\d+\s*章|\d+\.\d+)", text):
            headings.append((i, text))
    print("HEADINGS")
    for i, text in headings:
        if text.startswith("6.") or text.startswith("7.") or text.startswith("第 7") or text.startswith("第7"):
            print(i, text)

    print("\nCAPTIONS")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if re.match(r"^(图|表)\s*\d+[-－]\d+", text):
            print(i, text[:100])

    print("\nOVERSIZED IMAGE CHECK")
    for idx, shape in enumerate(doc.inline_shapes, start=1):
        width_in = shape.width / 914400
        height_in = shape.height / 914400
        print(idx, f"{width_in:.2f}x{height_in:.2f}in")


if __name__ == "__main__":
    main()
