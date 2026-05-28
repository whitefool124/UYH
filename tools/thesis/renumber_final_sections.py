from pathlib import Path

from docx import Document


DOCX = Path("毕业设计提交归档-曹逸天") / "3毕业设计说明书-曹逸天-学校模板图表补充精修版.docx"
MATERIAL = Path("论文材料") / "曹逸天-毕业设计说明书-学校模板Word重排版-实验增强版-图表补充精修版.docx"


def replace_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def main() -> None:
    doc = Document(DOCX)
    seen_yolo = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "6.8 YOLO 前置区域定位增强实验":
            seen_yolo = True
        elif seen_yolo and text == "6.8 补充数据统计与结论图表":
            replace_text(paragraph, "6.9 补充数据统计与结论图表")
            break
    doc.save(DOCX)
    doc.save(MATERIAL)
    print(DOCX)


if __name__ == "__main__":
    main()
