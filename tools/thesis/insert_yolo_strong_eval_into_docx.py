from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


DOCX = Path("毕业设计提交归档-曹逸天") / "3毕业设计说明书-曹逸天-学校模板图表补充精修版.docx"
MATERIAL = Path("论文材料") / "曹逸天-毕业设计说明书-学校模板Word重排版-实验增强版-图表补充精修版.docx"
SUMMARY = Path("论文材料") / "实验补充数据" / "yolo_roi_strong_summary.csv"
CSV_PATH = Path("论文材料") / "实验补充数据" / "yolo_roi_strong_eval.csv"
CHART_DIR = Path("毕业设计提交归档-曹逸天") / "论文图表与实验数据" / "ThesisAssets" / "experiment_charts_png"


def summary() -> dict[str, str]:
    with SUMMARY.open("r", encoding="utf-8-sig", newline="") as handle:
        return {row["metric"]: row["value"] for row in csv.DictReader(handle)}


def find_anchor(doc: Document):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("6.8 补充数据统计"):
            return paragraph
    raise ValueError("Cannot find 6.8 anchor")


def paragraph_before(anchor, text="", style=None):
    p = anchor.insert_paragraph_before(text)
    if style is not None:
        p.style = style
    return p


def body(anchor, text: str, style_source):
    p = paragraph_before(anchor, "", style_source.style)
    run = p.add_run(text)
    run.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Pt(24)
    return p


def image(anchor, path: Path, caption: str, note: str, style_source):
    p = paragraph_before(anchor, "")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(5.4))
    c = paragraph_before(anchor, caption, style_source.style)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.first_line_indent = None
    body(anchor, note, style_source)


def main() -> None:
    doc = Document(DOCX)
    anchor = find_anchor(doc)
    style_source = anchor
    s = summary()

    heading = paragraph_before(anchor, "6.8 YOLO 前置区域定位增强实验", style_source.style)
    heading.paragraph_format.first_line_indent = None

    body(
        anchor,
        "为进一步增强 YOLO 相关实验的证据力度，本文在外部桥接路线之外补充进行 YOLO 前置区域定位增强实验。实验不将 YOLO 作为动态手势分类器，而是检验其作为 person ROI 前置检测模块时，能否为后续 MediaPipe 关键点提取提供稳定候选区域。实验视频来自项目 bridge/samples 目录及其 ipn_real 子目录，共覆盖 9 个离线视频。",
        style_source,
    )
    body(
        anchor,
        f"实验使用 yolo11n.pt，置信度阈值设为 0.25，每段视频最多处理 120 帧。统计结果表明，9 个视频共处理 {s['frames']} 帧，YOLO 检出 person ROI 的帧数为 {s['yolo_person_frames']}，加权检出比例为 {s['weighted_person_ratio']}；平均 YOLO 置信度为 {s['avg_yolo_confidence']}，平均 person 区域面积占比为 {s['avg_person_area_ratio']}，平均处理耗时为 {s['avg_processing_ms']} ms，平均 FPS 为 {s['avg_fps']}。原始逐视频结果见 {CSV_PATH.as_posix()}。",
        style_source,
    )

    image(
        anchor,
        CHART_DIR / "yolo_roi_person_ratio.png",
        "图 6-7 YOLO 前置检测 person ROI 检出比例",
        "如图 6-7 所示，在本组离线样本中 YOLO 对 person ROI 的检出比例达到 1.0，说明其能够稳定完成前置区域定位任务。这一结果可以支撑将 YOLO 作为复杂背景下候选区域约束模块，而不是直接替代 MediaPipe 的手部关键点检测。",
        style_source,
    )
    image(
        anchor,
        CHART_DIR / "yolo_roi_speed_confidence.png",
        "图 6-8 YOLO 前置检测速度与置信度对比",
        "如图 6-8 所示，各视频平均置信度较高，平均处理速度约为 17 FPS。该速度低于 Mock 和 ExternalBridge 的游戏侧运行帧率，但作为离线回放、外部桥接验证或后续复杂背景增强实验仍具有可用性。",
        style_source,
    )
    body(
        anchor,
        "综合来看，YOLO 强化实验能够证明本文视觉链路具备前置区域定位扩展能力：在人物区域可稳定检出的情况下，系统可以先利用 YOLO 约束图像候选区域，再交由 MediaPipe 或其他关键点模块进行细粒度手势分析。但该实验同时表明，YOLO 当前的贡献边界仍是区域定位增强，不应被表述为独立动态手势识别模型成果。",
        style_source,
    )

    doc.save(DOCX)
    doc.save(MATERIAL)
    print(DOCX)
    print(MATERIAL)


if __name__ == "__main__":
    main()
