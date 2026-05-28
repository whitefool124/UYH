from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


DOCX = Path("毕业设计提交归档-曹逸天") / "3毕业设计说明书-曹逸天-学校模板图表补充精修版.docx"
MATERIAL = Path("论文材料") / "曹逸天-毕业设计说明书-学校模板Word重排版-实验增强版-图表补充精修版.docx"


def replace_para(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def insert_before(anchor, text: str = "", style=None):
    p = anchor.insert_paragraph_before(text)
    if style is not None:
        p.style = style
    return p


def body(anchor, text: str, style_source):
    p = insert_before(anchor, "", style_source.style)
    run = p.add_run(text)
    run.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Pt(24)
    return p


def find_para(doc: Document, predicate):
    for paragraph in doc.paragraphs:
        if predicate(paragraph.text.strip()):
            return paragraph
    raise ValueError("anchor not found")


def has_text(doc: Document, needle: str) -> bool:
    return any(needle in p.text for p in doc.paragraphs)


def add_variable_table(doc: Document) -> None:
    if has_text(doc, "表 4-5 动态手势识别算法变量说明"):
        return
    anchor = find_para(doc, lambda t: t.startswith("4.6"))
    style_source = anchor
    body(
        anchor,
        "为便于理解算法 4-1 中各参数的含义，本文将动态手势轨迹识别过程中使用的主要变量整理如表 4-5 所示。相关变量并不是固定绑定某一种摄像头或输入源，而是在统一手势帧和命令层中使用，因此可以同时服务于 Mock、Native MediaPipe 和 ExternalBridge 三类输入。",
        style_source,
    )
    caption = insert_before(anchor, "表 4-5 动态手势识别算法变量说明", style_source.style)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = None

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    headers = ["变量", "含义", "主要作用", "论文中的使用位置"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    rows = [
        ["W", "时间窗长度", "限定连续轨迹分析的时间范围，避免过旧帧影响当前判断", "滑动、响指、姿态转换识别"],
        ["H", "历史帧窗口", "保存近期 GestureFrame，用于计算位移、速度和状态变化", "动态轨迹计算"],
        ["C", "最小置信度阈值", "过滤低质量手部关键点或外部视觉帧", "输入有效性检查"],
        ["D", "最小位移阈值", "判断掌心或关键点轨迹是否达到动作触发条件", "滑动和身体位移动作"],
        ["R", "状态转换规则", "描述从一种手势状态到另一种手势状态的合法变化", "指向到握拳、响指等动作"],
        ["T", "冷却时间", "抑制短时间内重复触发同一动态命令", "命令过滤与游戏交互稳定性"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9)
    anchor._p.addprevious(table._tbl)


def add_real_camera_section(doc: Document) -> None:
    if has_text(doc, "6.10 真实摄像头演示验证与局限"):
        return
    anchor = find_para(doc, lambda t: t.startswith("第 7 章") or t.startswith("第7章"))
    style_source = anchor
    heading = insert_before(anchor, "6.10 真实摄像头演示验证与局限", style_source.style)
    heading.paragraph_format.first_line_indent = None
    body(
        anchor,
        "除离线回放和训练集抽样实验外，本文还通过真实摄像头对系统主要流程进行演示验证。验证内容包括开始菜单进入、玩法说明查看、摄像头校准、战斗场景运行、开发者实验室调试以及自定义手势验证等环节。相关截图已经归档在 ThesisAssets/screenshots 目录中，并在第 5 章作为系统运行效果证据使用。",
        style_source,
    )
    body(
        anchor,
        "真实摄像头验证的意义在于确认系统并非仅能处理离线样本，而是能够在实际输入链路下完成从视觉帧、手势命令到 Unity 游戏反馈的闭环。但该验证仍属于开发者演示与功能验证，不等同于多人、长时间、跨环境的用户实验。光照条件、背景复杂度、摄像头摆放角度、玩家手型差异和动作速度仍可能影响识别稳定性。",
        style_source,
    )
    body(
        anchor,
        "因此，本文将真实摄像头结果作为系统可运行性和工程闭环的补充证据，而将量化分析主要建立在 Jester-300 抽样、外部模板回放、YOLO ROI 检测和三输入模式性能记录之上。后续若要进一步提升结论强度，应补充多名用户、多环境和长时间交互测试，并统计误触发率、漏检率和用户主观体验指标。",
        style_source,
    )


def add_contribution_summary(doc: Document) -> None:
    if has_text(doc, "从贡献类型来看，本文成果可以归纳为方法、工程和实验三个层面"):
        return
    anchor = find_para(doc, lambda t: t.startswith("7.2"))
    style_source = anchor
    # Insert after the 7.2 heading by adding before the first following 7.3 if possible,
    # otherwise directly before 7.2 would be less useful.
    next_anchor = find_para(doc, lambda t: t.startswith("7.3"))
    body(
        next_anchor,
        "从贡献类型来看，本文成果可以归纳为方法、工程和实验三个层面。方法层面，本文围绕动态手势轨迹跟踪设计了基于时间窗、位移阈值、状态转换和冷却机制的命令识别方法，并通过变量表和伪代码明确其处理流程。工程层面，本文构建了 Mock、Native MediaPipe 和 ExternalBridge 三类输入源的统一抽象，使视觉层输出能够以 GestureFrame 和 GestureCommand 的形式进入 Unity 游戏逻辑。实验层面，本文通过三输入模式性能记录、Jester-300 抽样回放、YOLO ROI 前置检测和自定义手势增强采集分析，对系统可用性、实时性和局限性进行了验证。",
        style_source,
    )


def main() -> None:
    doc = Document(DOCX)
    add_variable_table(doc)
    add_real_camera_section(doc)
    add_contribution_summary(doc)
    doc.save(DOCX)
    doc.save(MATERIAL)
    print(DOCX)
    print(MATERIAL)


if __name__ == "__main__":
    main()
