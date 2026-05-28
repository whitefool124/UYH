from pathlib import Path
import re

from docx import Document


def find_docx(root: Path, pattern: str) -> Path:
    matches = list(root.rglob(pattern))
    if not matches:
        raise FileNotFoundError(pattern)
    return matches[0]


def extract(path: Path):
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    headings = []
    for i, text in enumerate(paragraphs):
        if (
            re.match(r"^第\s*\d+\s*章", text)
            or re.match(r"^\d+(\.\d+){0,2}\s+", text)
            or text in {"摘要", "Abstract", "参考文献", "致谢"}
        ):
            headings.append((i, text[:120]))
    refs = [t for t in paragraphs if re.match(r"^\[\d+\]", t)]
    fig_caps = [t for t in paragraphs if re.match(r"^图\s*\d+[-－]\d+", t)]
    table_caps = [t for t in paragraphs if re.match(r"^表\s*\d+[-－]\d+", t)]
    algo_hits = [t for t in paragraphs if any(k in t for k in ["算法", "伪代码", "输入：", "输出：", "阈值", "时间窗"])]
    return {
        "path": str(path),
        "paragraphs": len(paragraphs),
        "tables": len(doc.tables),
        "images": len(doc.inline_shapes),
        "headings": headings,
        "refs": len(refs),
        "fig_caps": fig_caps,
        "table_caps": table_caps,
        "algo_hits": algo_hits[:30],
    }


mine = extract(Path("毕业设计提交归档-曹逸天") / "3毕业设计说明书-曹逸天.docx")
example = extract(find_docx(Path("build-temp/example_zip"), "3*.docx"))

for name, data in [("MINE", mine), ("EXAMPLE", example)]:
    print(f"\n== {name} ==")
    print(data["path"])
    print("paragraphs", data["paragraphs"], "tables", data["tables"], "images", data["images"], "refs", data["refs"])
    print("-- headings --")
    for _, h in data["headings"][:80]:
        print(h)
    print("-- figures --")
    for h in data["fig_caps"][:30]:
        print(h)
    print("-- tables --")
    for h in data["table_caps"][:30]:
        print(h)
    print("-- algo hits --")
    for h in data["algo_hits"][:20]:
        print(h[:160])
