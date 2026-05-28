from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt


SOURCE = Path("毕业设计提交归档-曹逸天") / "3毕业设计说明书-曹逸天.docx"
TARGET = Path("毕业设计提交归档-曹逸天") / "3毕业设计说明书-曹逸天-内容风险补强版.docx"


def clone_style(dst, src):
    dst.style = src.style
    dst.paragraph_format.left_indent = src.paragraph_format.left_indent
    dst.paragraph_format.first_line_indent = src.paragraph_format.first_line_indent
    dst.paragraph_format.space_before = src.paragraph_format.space_before
    dst.paragraph_format.space_after = src.paragraph_format.space_after
    dst.paragraph_format.line_spacing = src.paragraph_format.line_spacing
    dst.alignment = src.alignment


def insert_paragraph_before(paragraph, text="", style_from=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    new_para._element = new_p
    if style_from is not None:
        clone_style(new_para, style_from)
    if text:
        run = new_para.add_run(text)
        if style_from is not None and style_from.runs:
            run.font.name = style_from.runs[0].font.name
            run.font.size = style_from.runs[0].font.size
    return new_para


def insert_table_before(paragraph, rows, cols):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Pt(450))
    paragraph._p.addprevious(table._tbl)
    return table


def find_para(doc, text):
    for para in doc.paragraphs:
        if para.text.strip() == text:
            return para
    raise ValueError(f"paragraph not found: {text}")


def find_para_prefix(doc, prefix):
    for para in doc.paragraphs:
        if para.text.strip().startswith(prefix):
            return para
    raise ValueError(f"paragraph prefix not found: {prefix}")


def previous_paragraph_style(doc, anchor):
    for index, para in enumerate(doc.paragraphs):
        if para._p is anchor._p:
            return doc.paragraphs[max(0, index - 1)]
    return doc.paragraphs[0]


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_method_comparison(doc):
    anchor = find_para_prefix(doc, "2.5")
    body_style = previous_paragraph_style(doc, anchor)

    insert_paragraph_before(anchor, "", body_style)
    intro = (
        "为进一步明确本文方法与现有研究之间的关系，表 2-1 从输入特征、适用场景、"
        "主要优势和局限性等角度对常见动态手势识别方案进行比较。可以看出，深度时序模型"
        "在大规模数据和复杂动作识别方面具有优势，但对数据规模、训练成本和部署环境要求较高；"
        "关键点轨迹规则方法虽然表达能力有限，但实时性、可解释性和工程接入成本更适合本文的"
        "体感游戏原型验证场景。因此，本文没有将研究目标设定为训练通用手势分类模型，而是围绕"
        "普通 RGB 摄像头条件下的轨迹跟踪、命令抽象和 Unity 游戏闭环进行改进。"
    )
    insert_paragraph_before(anchor, intro, body_style)
    cap = insert_paragraph_before(anchor, "表 2-1 现有动态手势识别方法对比", body_style)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    data = [
        ["方法类别", "代表思路", "优势", "不足", "本文采用或改进点"],
        ["传统视觉方法", "肤色分割、背景差分、轮廓特征", "实现简单，计算量较低", "受光照、背景和肤色差异影响较大", "作为早期方法背景，不作为主线实现"],
        ["目标检测方法", "YOLO 等检测手部或人体区域", "复杂背景下定位能力较强", "仅给出区域，难以直接表达动态轨迹", "作为外部视觉桥接和前置区域定位增强方向"],
        ["关键点规则方法", "MediaPipe 关键点、时间窗、位移阈值", "实时性好，可解释，便于接入游戏逻辑", "依赖阈值，面对个体差异时泛化有限", "作为本文核心方法，结合命令抽象、冷却机制和状态转换"],
        ["模板匹配方法", "DTW 或轨迹模板相似度", "可处理不同长度的动作序列", "模板质量和边界定义影响较大", "用于自定义手势模板验证，并分析少样本风险"],
        ["深度时序模型", "LSTM、GRU、Transformer", "适合复杂动态手势和大规模数据", "训练成本高，数据需求大，可解释性较弱", "作为后续改进计划，与规则方法形成对比"],
    ]
    table = insert_table_before(anchor, len(data), len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for r, row in enumerate(data):
        for c, value in enumerate(row):
            set_cell_text(table.cell(r, c), value, bold=(r == 0))
    insert_paragraph_before(anchor, "", body_style)


def add_algorithm_pseudocode(doc):
    anchor = find_para_prefix(doc, "4.6")
    body_style = previous_paragraph_style(doc, anchor)
    title = insert_paragraph_before(anchor, "算法 4-1 基于时间窗的动态手势轨迹命令识别算法", body_style)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines = [
        "输入：连续手势帧序列 F、时间窗长度 W、最小置信度阈值 C、位移阈值 D、状态转换规则 R、冷却时间 T。",
        "输出：动态手势命令 GestureCommand，若当前窗口不满足触发条件则输出空命令。",
        "（1）接收当前手势帧 f_t，检查时间戳、置信度和关键点有效性；若 f_t 低于 C，则仅更新丢失状态并返回空命令。",
        "（2）将 f_t 写入历史窗口 H，并删除时间跨度超过 W 的旧帧，保证 H 只保留近期连续轨迹。",
        "（3）根据掌心、指尖和手腕关键点计算归一化位移、方向、速度以及关键指尖距离变化。",
        "（4）若系统仍处于同类命令冷却时间 T 内，则抑制重复触发，返回空命令。",
        "（5）依次判断滑动、响指、指向到握拳等状态转换规则 R：当位移方向、持续时间、手型状态和最小幅度均满足规则时，生成候选命令。",
        "（6）结合 GestureCommandHistory 检查候选命令是否与当前游戏状态和最近命令冲突；若冲突，则丢弃候选命令。",
        "（7）输出通过过滤的 GestureCommand，并记录触发时间、输入来源、置信度和动作类型；若无候选命令通过过滤，则输出空命令。",
    ]
    for line in lines:
        p = insert_paragraph_before(anchor, line, body_style)
        for run in p.runs:
            run.font.size = Pt(10.5)
    note = (
        "该算法体现了本文方法的核心思想：视觉层不直接驱动游戏对象，而是先在连续帧窗口中形成"
        "可解释的动态命令，再由玩法层根据当前状态消费命令。与端到端分类模型相比，该方法牺牲了"
        "一部分复杂动作泛化能力，但换取了实时性、可调试性和毕业设计原型中的工程可复现性。"
    )
    insert_paragraph_before(anchor, note, body_style)
    insert_paragraph_before(anchor, "", body_style)


def add_experiment_questions(doc):
    anchor = find_para_prefix(doc, "6.1")
    body_style = previous_paragraph_style(doc, anchor)
    text = (
        "本章实验并不将目标设定为证明系统已经达到通用手势识别模型的长期用户级准确率，"
        "而是围绕毕业设计系统的工程闭环和方法边界展开验证。具体而言，本章重点回答四个问题："
        "第一，系统是否能够完成从视觉输入、手势命令抽象到 Unity 游戏反馈的完整链路；第二，"
        "Mock、Native MediaPipe 和 ExternalBridge 三种输入模式在实时性指标上是否具备可演示性；"
        "第三，基于时间窗和阈值的动态轨迹识别在离线回放样本中是否能够稳定触发目标动作；第四，"
        "少样本自定义动态手势模板在真实摄像头交互中会暴露哪些边界问题。通过上述问题划分，"
        "本文将功能验证、性能记录、离线回放和局限性分析区分开来，避免将工程演示结果扩大解释为"
        "大规模真实用户实验结论。"
    )
    insert_paragraph_before(anchor, text, body_style)
    insert_paragraph_before(anchor, "", body_style)


def main():
    doc = Document(SOURCE)
    add_method_comparison(doc)
    add_algorithm_pseudocode(doc)
    add_experiment_questions(doc)
    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
