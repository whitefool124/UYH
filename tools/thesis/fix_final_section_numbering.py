from pathlib import Path

from docx import Document


DOCX = Path("毕业设计提交归档-曹逸天") / "3毕业设计说明书-曹逸天-学校模板图表补充精修版.docx"
MATERIAL = Path("论文材料") / "曹逸天-毕业设计说明书-学校模板Word重排版-实验增强版-图表补充精修版.docx"

REPLACEMENTS = {
    "6.8 YOLO 前置区域定位增强实验": "6.9 YOLO 前置区域定位增强实验",
    "6.9 补充数据统计与结论图表": "6.10 补充数据统计与结论图表",
    "6.10 真实摄像头演示验证与局限": "6.11 真实摄像头演示验证与局限",
}


def replace_text(paragraph, text):
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def main():
    doc = Document(DOCX)
    changed = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if text in REPLACEMENTS:
            replace_text(p, REPLACEMENTS[text])
            changed += 1
    doc.save(DOCX)
    doc.save(MATERIAL)
    print("changed", changed)


if __name__ == "__main__":
    main()
