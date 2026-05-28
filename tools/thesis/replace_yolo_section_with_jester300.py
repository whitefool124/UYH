from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


DOCX = Path("毕业设计提交归档-曹逸天") / "3毕业设计说明书-曹逸天-学校模板图表补充精修版.docx"
MATERIAL = Path("论文材料") / "曹逸天-毕业设计说明书-学校模板Word重排版-实验增强版-图表补充精修版.docx"
SUMMARY = Path("论文材料") / "实验补充数据" / "yolo_roi_jester300_summary.csv"
CSV_PATH = Path("论文材料") / "实验补充数据" / "yolo_roi_jester300_eval.csv"
CHART_DIR = Path("毕业设计提交归档-曹逸天") / "论文图表与实验数据" / "ThesisAssets" / "experiment_charts_png"


def read_summary() -> dict[str, str]:
    with SUMMARY.open("r", encoding="utf-8-sig", newline="") as handle:
        return {row["metric"]: row["value"] for row in csv.DictReader(handle)}


def replace_para(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def find_range(doc: Document) -> tuple[int, int]:
    start = None
    end = None
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == "6.8 YOLO 前置区域定位增强实验":
            start = idx
        elif start is not None and text.startswith("6.9 "):
            end = idx
            break
    if start is None or end is None:
        raise ValueError("Cannot locate YOLO section")
    return start, end


def insert_before(anchor, text="", style=None):
    p = anchor.insert_paragraph_before(text)
    if style is not None:
        p.style = style
    return p


def body(anchor, text: str, style_source):
    p = insert_before(anchor, "", style_source.style)
    run = p.add_run(text)
    run.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Pt(24)
    return p


def image(anchor, path: Path, caption: str, note: str, style_source):
    p = insert_before(anchor, "")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(5.4))
    c = insert_before(anchor, caption, style_source.style)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.first_line_indent = None
    body(anchor, note, style_source)


def main() -> None:
    doc = Document(DOCX)
    start, end = find_range(doc)
    anchor = doc.paragraphs[end]
    style_source = doc.paragraphs[start]
    # Delete old section from bottom to top.
    for idx in range(end - 1, start - 1, -1):
        delete_paragraph(doc.paragraphs[idx])

    s = read_summary()
    heading = insert_before(anchor, "6.8 YOLO 前置区域定位增强实验", style_source.style)
    heading.paragraph_format.first_line_indent = None
    body(
        anchor,
        "为进一步增强 YOLO 相关实验的证据力度，本文基于 Jester-300 抽样训练集补充进行 YOLO 前置区域定位增强实验。实验不将 YOLO 作为动态手势分类器，而是检验其作为 person ROI 前置检测模块时，能否为后续 MediaPipe 关键点提取提供稳定候选区域。Jester-300 来自大规模训练集抽样帧目录，共包含 300 个 clip，每个 clip 均匀抽取最多 8 帧进行检测。",
        style_source,
    )
    body(
        anchor,
        f"实验使用 yolo11n.pt，置信度阈值设为 0.25，共处理 {s['clips']} 个 clip、{s['frames']} 帧。YOLO 检出 person ROI 的帧数为 {s['yolo_person_frames']}，加权帧检出率为 {s['weighted_person_ratio']}；阳性 clip 数量为 {s['positive_clips']}，clip 阳性率为 {s['positive_clip_rate']}；平均 clip 检出率为 {s['avg_clip_person_ratio']}。在阳性帧中，平均 YOLO 置信度为 {s['avg_yolo_confidence_positive']}，平均 person 区域面积占比为 {s['avg_person_area_ratio_positive']}。平均处理耗时为 {s['avg_processing_ms']} ms，平均 FPS 为 {s['avg_fps']}。原始逐 clip 结果见 {CSV_PATH.as_posix()}。",
        style_source,
    )
    image(
        anchor,
        CHART_DIR / "yolo_jester300_summary_rates.png",
        "图 6-7 Jester-300 YOLO ROI 总体检出指标",
        "如图 6-7 所示，在 300 个训练集抽样 clip 上，YOLO 前置检测取得较高的帧检出率和 clip 阳性率，说明其可以稳定承担人物区域定位任务，为后续关键点检测提供候选区域约束。",
        style_source,
    )
    image(
        anchor,
        CHART_DIR / "yolo_jester300_clip_ratio_distribution.png",
        "图 6-8 Jester-300 YOLO ROI clip 检出比例分布",
        "如图 6-8 所示，大多数 clip 的 person ROI 检出比例集中在高区间，说明 YOLO 对训练集抽样帧具有较好的区域定位覆盖能力。该结果强化了本文外部视觉桥接链路的可扩展性证据。",
        style_source,
    )
    body(
        anchor,
        "综合来看，扩大后的 YOLO 实验能够更有力地证明本文视觉链路具备前置区域定位扩展能力：在训练集抽样规模扩大到 300 个 clip 后，YOLO 仍能保持较高 ROI 检出比例。但该实验的贡献边界仍是区域定位增强，后续若要形成独立手势识别模型，还需要进一步进行手部框标注、类别监督训练和真实摄像头用户测试。",
        style_source,
    )

    doc.save(DOCX)
    doc.save(MATERIAL)
    print(DOCX)
    print(MATERIAL)


if __name__ == "__main__":
    main()
