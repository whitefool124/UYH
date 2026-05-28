from __future__ import annotations

import re
import os
from collections import Counter
from pathlib import Path

from docx import Document


DOCX = Path(os.environ.get("FINAL_DOCX", "毕业设计提交归档-曹逸天/3毕业设计说明书-曹逸天-学校模板图表补充版.docx"))


def main() -> None:
    doc = Document(DOCX)
    paragraphs = [(i, p.text.strip()) for i, p in enumerate(doc.paragraphs) if p.text.strip()]
    all_text = "\n".join(text for _, text in paragraphs)

    print("DOCX", DOCX)
    print("paragraphs", len(doc.paragraphs), "nonempty", len(paragraphs), "tables", len(doc.tables), "images", len(doc.inline_shapes))

    patterns = {
        "todo_or_placeholder": r"待填|可选|\[ \]|TODO|FIXME|未完成|回填入口|可选补充",
        "template_warning": r"注意：本论文|抄袭本论文|仅供参考",
        "source_code_marker": r"```|public\s+|private\s+|protected\s+|class\s+\w+|void\s+\w+\s*\(",
        "overclaim_yolo": r"训练了.*YOLO|高精度\s*YOLO|YOLO.*分类器成果|完整\s*YOLO.*训练",
        "bad_encoding_marker": r"�|闈|绗|鍥|琛|鏂|妯",
    }
    for name, pattern in patterns.items():
        hits = [(i, text) for i, text in paragraphs if re.search(pattern, text, flags=re.I)]
        print(f"\n== {name}: {len(hits)} ==")
        for i, text in hits[:20]:
            print(i, text[:180])

    fig_caps = [(i, text) for i, text in paragraphs if re.match(r"^图\s*\d+[-－]\d+", text)]
    table_caps = [(i, text) for i, text in paragraphs if re.match(r"^表\s*\d+[-－]\d+", text)]
    refs = [(i, text) for i, text in paragraphs if re.match(r"^\[\d+\]", text)]

    print(f"\nfig_caps {len(fig_caps)}")
    fig_nums = []
    for i, text in fig_caps:
        m = re.match(r"^图\s*(\d+[-－]\d+)", text)
        fig_nums.append(m.group(1).replace("－", "-"))
    for num, count in Counter(fig_nums).items():
        if count > 1:
            print("DUP_FIG", num, count)

    print(f"\ntable_caps {len(table_caps)}")
    table_nums = []
    for i, text in table_caps:
        m = re.match(r"^表\s*(\d+[-－]\d+)", text)
        table_nums.append(m.group(1).replace("－", "-"))
    for num, count in Counter(table_nums).items():
        if count > 1:
            print("DUP_TABLE", num, count)

    print(f"\nrefs {len(refs)}")
    ref_nums = []
    for _, text in refs:
        m = re.match(r"^\[(\d+)\]", text)
        ref_nums.append(int(m.group(1)))
    if ref_nums:
        missing = sorted(set(range(1, max(ref_nums) + 1)) - set(ref_nums))
        dupes = [num for num, count in Counter(ref_nums).items() if count > 1]
        print("ref range", min(ref_nums), max(ref_nums), "missing", missing, "dupes", dupes)

    suspicious_refs = []
    for i, text in refs:
        if "[J/OL]" in text or "[EB/OL]" in text:
            if "http" not in text and "DOI" not in text and "doi" not in text and "arXiv" not in text:
                suspicious_refs.append((i, text))
    print(f"\nsuspicious_refs {len(suspicious_refs)}")
    for i, text in suspicious_refs[:20]:
        print(i, text)

    chapter7 = [i for i, text in paragraphs if text.startswith("第 7 章") or text.startswith("第7章")]
    section68 = [i for i, text in paragraphs if text.startswith("6.8")]
    print("\nsection positions", "6.8", section68[:5], "chapter7", chapter7[:5])


if __name__ == "__main__":
    main()
