from pathlib import Path

from docx import Document


SOURCE = Path("毕业设计提交归档-曹逸天") / "3毕业设计说明书-曹逸天-学校模板图表补充精修版.docx"
MATERIAL = Path("论文材料") / "曹逸天-毕业设计说明书-学校模板Word重排版-实验增强版-图表补充精修版.docx"


REPLACEMENTS = {
    "当前尚未完成大规模 YOLO 手势模型训练，因此论文中不能声称系统训练了高精度 YOLO 手势分类模型；现有 YOLO 数据只适合作为外部桥接前端和区域定位增强的可行性证据。": "当前 YOLO 相关工作定位为外部桥接前端和区域定位增强的可行性验证，尚未形成可作为核心成果的独立手势分类模型。",
    "YOLO+MediaPipe 当前应表述为外部桥接可行性和后续优化方向，不能表述为已完成高精度 YOLO 动态手势分类器。": "YOLO+MediaPipe 当前应表述为外部桥接可行性和后续优化方向，不应扩展解释为独立动态手势分类模型成果。",
    "本实验的目的不是训练一个最终可部署的手势分类模型，而是回答在给定模板候选视频和离线回放条件下，规则模板是否具有最小可验证稳定性这一问题。": "本实验的目的不是构建最终可部署的手势分类模型，而是回答在给定模板候选视频和离线回放条件下，规则模板是否具有最小可验证稳定性这一问题。",
    "本实验的目的不是构建最终可部署的手势分类模型，而是回答在给定模板候选视频和离线回放条件下，规则模板是否具有最小可验证稳定性这一问题。": "本实验的目的在于回答：给定模板候选视频和离线回放条件下，规则模板是否具有最小可验证稳定性。",
}


def replace_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def main() -> None:
    doc = Document(SOURCE)
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text
        new_text = text
        for old, new in REPLACEMENTS.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            replace_paragraph_text(paragraph, new_text)
            changed += 1
        elif "最终可部署" in text and "规则模板" in text:
            replace_paragraph_text(
                paragraph,
                "本组实验用于回答在给定模板候选视频和离线回放条件下，规则模板是否具有最小可验证稳定性。本文关注目录生成、采集 clip 数、回放验证 clip 数、模板回放准确率、验证 clip 正确数量、手部检出比例和三种输入模式性能指标。Jester-120 与 Jester-300 的对比用于观察样本规模和标签分布变化对规则模板筛选的影响。",
            )
            changed += 1
    doc.save(SOURCE)
    doc.save(MATERIAL)
    print("changed", changed)
    print(SOURCE)
    print(MATERIAL)


if __name__ == "__main__":
    main()
