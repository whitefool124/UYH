from __future__ import annotations

import re
from pathlib import Path

from docx import Document


SOURCE = Path("毕业设计提交归档-曹逸天") / "3毕业设计说明书-曹逸天-学校模板图表补充版.docx"
TARGET = Path("毕业设计提交归档-曹逸天") / "3毕业设计说明书-曹逸天-学校模板图表补充精修版.docx"
MATERIAL_TARGET = Path("论文材料") / "曹逸天-毕业设计说明书-学校模板Word重排版-实验增强版-图表补充精修版.docx"


REF_FIXES = {
    "[2]": "[2] Zhang F, Bazarevsky V, Vakunov A, et al. MediaPipe Hands: On-device Real-time Hand Tracking[EB/OL]. Google AI Blog, 2019. https://developers.googleblog.com/2019/08/on-device-real-time-hand-tracking-with.html.",
    "[11]": "[11] 动态手势理解与交互综述[J]. 软件学报, 2021.",
    "[19]": "[19] Zheng Y, et al. Gesture Recognition Based on MediaPipe for Computer Game Control[J]. 2023.",
    "[24]": "[24] Multimodal Vision-based Human Activity Recognition Using Deep Learning: A Review[J]. 2024.",
}


def replace_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def main() -> None:
    doc = Document(SOURCE)
    caption_seen: set[str] = set()
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        # Some explanatory paragraphs start with "图 x-y ...", which makes them
        # indistinguishable from real figure captions in audits. Keep the content
        # but rephrase them as body text.
        match = re.match(r"^(图\s*\d+[-－]\d+)\s+(.+)$", text)
        if match:
            fig_no = match.group(1).replace("－", "-")
            rest = match.group(2)
            if fig_no in caption_seen or len(text) > 42:
                replace_paragraph_text(paragraph, f"如{fig_no}所示，{rest}")
            else:
                caption_seen.add(fig_no)

        stripped = paragraph.text.strip()
        for prefix, replacement in REF_FIXES.items():
            if stripped.startswith(prefix):
                replace_paragraph_text(paragraph, replacement)
                break

    doc.save(TARGET)
    doc.save(MATERIAL_TARGET)
    print(TARGET)
    print(MATERIAL_TARGET)


if __name__ == "__main__":
    main()
