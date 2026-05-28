from pathlib import Path

from docx import Document


DOCX = Path("毕业设计提交归档-曹逸天") / "3毕业设计说明书-曹逸天-学校模板图表补充精修版.docx"
MATERIAL = Path("论文材料") / "曹逸天-毕业设计说明书-学校模板Word重排版-实验增强版-图表补充精修版.docx"

REPL = {
    "图 6-7 Jester-300 YOLO ROI 总体检出指标": "图 6-9 Jester-300 YOLO ROI 总体检出指标",
    "图 6-8 Jester-300 YOLO ROI clip 检出比例分布": "图 6-10 Jester-300 YOLO ROI clip 检出比例分布",
    "如图 6-7 所示": "如图 6-9 所示",
    "如图 6-8 所示": "如图 6-10 所示",
}


def set_text(p, text):
    if not p.runs:
        p.add_run(text)
        return
    p.runs[0].text = text
    for run in p.runs[1:]:
        run.text = ""


def main():
    doc = Document(DOCX)
    changed = 0
    for p in doc.paragraphs:
        text = p.text
        new = text
        for old, repl in REPL.items():
            new = new.replace(old, repl)
        if new != text:
            set_text(p, new)
            changed += 1
    doc.save(DOCX)
    doc.save(MATERIAL)
    print("changed", changed)


if __name__ == "__main__":
    main()
