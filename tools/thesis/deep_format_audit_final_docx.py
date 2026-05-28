from __future__ import annotations

import os
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


DOCX = Path(os.environ.get("FINAL_DOCX", "毕业设计提交归档-曹逸天/3毕业设计说明书-曹逸天-学校模板图表补充精修版.docx"))


FIG_RE = re.compile(r"^图\s*(\d+)-(\d+)\s+(.+)")
TAB_RE = re.compile(r"^表\s*(\d+)-(\d+)\s+(.+)")
SECTION_RE = re.compile(r"^(\d+)\.(\d+)(?:\.(\d+))?\s+")


def iter_blocks(doc: Document):
    body = doc.element.body
    para_map = {p._p: p for p in doc.paragraphs}
    table_map = {t._tbl: t for t in doc.tables}
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            p = para_map.get(child)
            if p is not None:
                yield "p", p
        elif child.tag == qn("w:tbl"):
            t = table_map.get(child)
            if t is not None:
                yield "tbl", t


def paragraph_has_image(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing | .//w:pict"))


def norm(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip())


def main() -> None:
    doc = Document(DOCX)
    issues: list[str] = []
    blocks = list(iter_blocks(doc))

    headings = []
    figures = []
    tables = []
    image_blocks = []
    for idx, (kind, obj) in enumerate(blocks):
        if kind == "p":
            text = norm(obj.text)
            if SECTION_RE.match(text) or text.startswith("第 ") or text.startswith("附录"):
                headings.append((idx, text))
            m = FIG_RE.match(text)
            if m:
                figures.append((idx, int(m.group(1)), int(m.group(2)), text))
            m = TAB_RE.match(text)
            if m:
                tables.append((idx, int(m.group(1)), int(m.group(2)), text))
            if paragraph_has_image(obj):
                image_blocks.append(idx)

    # Chapter 6 figure order in the body must be monotonic.
    chapter6 = [(idx, no, text) for idx, ch, no, text in figures if ch == 6 and idx < 390]
    nums = [no for _, no, _ in chapter6]
    if nums != sorted(nums):
        issues.append(f"Chapter 6 figure numbers are not monotonic: {nums}")
    if len(nums) != len(set(nums)):
        issues.append(f"Chapter 6 duplicate figure numbers: {nums}")

    # Captions should sit near their visual/table.
    for idx, ch, no, text in figures:
        prev_has_image = any(j in image_blocks for j in range(max(0, idx - 2), idx))
        if not prev_has_image and idx < 390:
            issues.append(f"Figure caption may not follow image near block {idx}: {text}")

    table_block_indices = [i for i, (kind, _) in enumerate(blocks) if kind == "tbl"]
    for idx, ch, no, text in tables:
        next_has_table = any(j in table_block_indices for j in range(idx + 1, min(len(blocks), idx + 4)))
        # Body table captions should precede tables; ignore appendix inventory if any.
        if idx < 390 and not next_has_table:
            issues.append(f"Table caption may not precede table near block {idx}: {text}")

    oversized = []
    for i, shape in enumerate(doc.inline_shapes, 1):
        width_in = shape.width / 914400
        height_in = shape.height / 914400
        area_ratio = (width_in * height_in) / (8.27 * 11.69)
        if area_ratio > 0.50 or height_in > 5.84:
            oversized.append((i, width_in, height_in, area_ratio))
    for i, w, h, r in oversized:
        issues.append(f"Image {i} oversized: {w:.2f} x {h:.2f} in, page area ratio {r:.2%}")

    # Risky text patterns.
    all_text = "\n".join(p.text for p in doc.paragraphs)
    patterns = {
        "source code fence": "```",
        "TODO": "TODO",
        "placeholder": "待补",
        "bad encoding marker": "�",
        "overclaim YOLO classifier": "YOLO 动态手势分类器",
    }
    for label, token in patterns.items():
        if token in all_text:
            issues.append(f"Risky token found: {label} / {token}")

    refs = []
    for p in doc.paragraphs:
        m = re.match(r"^\[(\d+)\]", p.text.strip())
        if m:
            refs.append(int(m.group(1)))
    if refs:
        expected = list(range(min(refs), max(refs) + 1))
        if refs != expected:
            issues.append(f"Reference list not continuous/orderly: {refs}")

    print(f"DOCX: {DOCX}")
    print(f"blocks={len(blocks)} paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} images={len(doc.inline_shapes)}")
    print(f"chapter6_body_figures={nums}")
    print(f"references={refs[0] if refs else 0}-{refs[-1] if refs else 0} count={len(refs)}")
    print("issues:")
    if issues:
        for issue in issues:
            print(f"- {issue}")
    else:
        print("- none")


if __name__ == "__main__":
    main()
