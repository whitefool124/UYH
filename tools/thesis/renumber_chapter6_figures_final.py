from pathlib import Path

from docx import Document


DOCX = Path("毕业设计提交归档-曹逸天") / "3毕业设计说明书-曹逸天-学校模板图表补充精修版.docx"
MATERIAL = Path("论文材料") / "曹逸天-毕业设计说明书-学校模板Word重排版-实验增强版-图表补充精修版.docx"


START = "6.9 YOLO 前置区域定位增强实验"
END = "6.11 真实摄像头演示验证与局限"


REPLACEMENTS = {
    "图 6-9 Jester-300 YOLO ROI 总体检出指标": "图 6-2 Jester-300 YOLO ROI 总体检出指标",
    "图 6-10 Jester-300 YOLO ROI clip 检出比例分布": "图 6-3 Jester-300 YOLO ROI clip 检出比例分布",
    "图 6-2 三种输入模式实时性能对比": "图 6-4 三种输入模式实时性能对比",
    "图 6-3 YOLO+MediaPipe 外部桥接有效检出分布": "图 6-5 YOLO+MediaPipe 外部桥接有效检出分布",
    "图 6-4 Jester 样本筛选拒绝原因统计": "图 6-6 Jester 样本筛选拒绝原因统计",
    "图 6-5 外部模板回放动态手势识别 F1": "图 6-7 外部模板回放动态手势识别 F1",
    "图 6-6 手指级变化与掌心轨迹比例示例": "图 6-8 手指级变化与掌心轨迹比例示例",
    "如图 6-9 所示": "如图 6-2 所示",
    "如图 6-10 所示": "如图 6-3 所示",
    "如图 6-2所示": "如图 6-4 所示",
    "如图 6-3所示": "如图 6-5 所示",
    "如图 6-4所示": "如图 6-6 所示",
    "如图 6-5所示": "如图 6-7 所示",
    "如图 6-6所示": "如图 6-8 所示",
}


def set_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def main() -> None:
    doc = Document(DOCX)
    in_scope = False
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text
        stripped = text.strip()
        if stripped == START:
            in_scope = True
        elif stripped == END:
            in_scope = False
        if not in_scope:
            continue
        new = text
        for old, replacement in REPLACEMENTS.items():
            new = new.replace(old, replacement)
        if new != text:
            set_text(paragraph, new)
            changed += 1

    doc.save(DOCX)
    doc.save(MATERIAL)
    print(f"changed paragraphs: {changed}")
    print(DOCX)
    print(MATERIAL)


if __name__ == "__main__":
    main()
