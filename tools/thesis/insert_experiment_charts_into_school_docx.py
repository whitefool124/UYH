from __future__ import annotations

import csv
import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path("论文材料") / "实验补充数据"
SOURCE = Path("论文材料") / "曹逸天-毕业设计说明书-学校模板Word重排版-实验增强版.docx"
TARGET = Path("论文材料") / "曹逸天-毕业设计说明书-学校模板Word重排版-实验增强版-图表补充版.docx"
ARCHIVE_COPY = Path("毕业设计提交归档-曹逸天") / "3毕业设计说明书-曹逸天-学校模板图表补充版.docx"
ASSET_DIR = Path("毕业设计提交归档-曹逸天") / "论文图表与实验数据" / "ThesisAssets" / "experiment_charts_png"
FONT = Path("C:/Windows/Fonts/msyh.ttc")
BOLD_FONT = Path("C:/Windows/Fonts/msyhbd.ttc")


def read_csv(name: str) -> list[dict[str, str]]:
    with (ROOT / name).open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(BOLD_FONT if bold else FONT), size)


def text_center(draw: ImageDraw.ImageDraw, xy: tuple[float, float], text: str, fnt, fill="#22313f") -> None:
    box = draw.textbbox((0, 0), text, font=fnt)
    draw.text((xy[0] - (box[2] - box[0]) / 2, xy[1] - (box[3] - box[1]) / 2), text, font=fnt, fill=fill)


def draw_bar_chart(path: Path, title: str, labels: list[str], values: list[float], ylabel: str, color: str) -> None:
    width, height = 1500, 900
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(42, True)
    label_font = font(25)
    small_font = font(22)
    text_center(draw, (width / 2, 65), title, title_font)

    left, top, right, bottom = 170, 150, 80, 155
    chart_w = width - left - right
    chart_h = height - top - bottom
    max_v = max(values) * 1.18 if values else 1
    max_v = max(max_v, 1)

    for i in range(5):
        y = top + chart_h - chart_h * i / 4
        value = max_v * i / 4
        draw.line([(left, y), (width - right, y)], fill="#d9e0e8", width=2)
        draw.text((left - 95, y - 15), f"{value:.1f}", font=small_font, fill="#52606d")

    draw.line([(left, top), (left, top + chart_h)], fill="#8c99a6", width=3)
    draw.line([(left, top + chart_h), (width - right, top + chart_h)], fill="#8c99a6", width=3)
    draw.text((34, top + chart_h / 2 - 18), ylabel, font=label_font, fill="#22313f")

    group_w = chart_w / len(values)
    bar_w = group_w * 0.55
    for index, (label, value) in enumerate(zip(labels, values)):
        x = left + group_w * index + group_w / 2 - bar_w / 2
        h = chart_h * value / max_v
        y = top + chart_h - h
        draw.rounded_rectangle([x, y, x + bar_w, top + chart_h], radius=8, fill=color)
        text_center(draw, (x + bar_w / 2, y - 28), f"{value:.3g}", small_font, "#22313f")
        text_center(draw, (x + bar_w / 2, top + chart_h + 45), label, small_font, "#22313f")

    img.save(path)


def draw_performance(path: Path) -> None:
    data = read_csv("input_mode_performance_summary.csv")
    modes = [r["mode"] for r in data]
    fps = [float(r["average_fps_mean"]) for r in data]
    p95 = [float(r["p95_frame_ms_mean"]) for r in data]

    width, height = 1500, 900
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(42, True)
    small_font = font(22)
    text_center(draw, (width / 2, 65), "三种输入模式实时性能对比", title_font)

    left, top, right, bottom = 155, 155, 85, 160
    chart_w = width - left - right
    chart_h = height - top - bottom
    max_v = max(max(fps), max(p95)) * 1.25
    for i in range(5):
        y = top + chart_h - chart_h * i / 4
        value = max_v * i / 4
        draw.line([(left, y), (width - right, y)], fill="#d9e0e8", width=2)
        draw.text((left - 90, y - 15), f"{value:.1f}", font=small_font, fill="#52606d")
    draw.line([(left, top + chart_h), (width - right, top + chart_h)], fill="#8c99a6", width=3)
    draw.rectangle([width - 350, 110, width - 325, 135], fill="#2f6f9f")
    draw.text((width - 315, 106), "平均 FPS", font=small_font, fill="#22313f")
    draw.rectangle([width - 200, 110, width - 175, 135], fill="#d08a2e")
    draw.text((width - 165, 106), "P95(ms)", font=small_font, fill="#22313f")

    group_w = chart_w / len(modes)
    bar_w = group_w * 0.22
    for i, mode in enumerate(modes):
        cx = left + group_w * i + group_w / 2
        for dx, value, color in [(-bar_w * 0.65, fps[i], "#2f6f9f"), (bar_w * 0.65, p95[i], "#d08a2e")]:
            h = chart_h * value / max_v
            x = cx + dx - bar_w / 2
            y = top + chart_h - h
            draw.rounded_rectangle([x, y, x + bar_w, top + chart_h], radius=7, fill=color)
            text_center(draw, (x + bar_w / 2, y - 25), f"{value:.1f}", small_font)
        text_center(draw, (cx, top + chart_h + 45), mode, small_font)
    img.save(path)


def draw_yolo(path: Path) -> None:
    summary = {r["metric"]: float(r["value"]) for r in read_csv("yolo_bridge_summary.csv")}
    positive = int(summary["hand_positive_videos"])
    zero = int(summary["zero_hand_videos"])
    total = positive + zero
    width, height = 1300, 760
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(40, True)
    label_font = font(28)
    small_font = font(24)
    text_center(draw, (width / 2, 65), "YOLO+MediaPipe 外部桥接有效检出分布", title_font)

    cx, cy, r = 420, 410, 220
    draw.pieslice([cx - r, cy - r, cx + r, cy + r], start=0, end=360, fill="#b35f5f")
    draw.pieslice([cx - r, cy - r, cx + r, cy + r], start=-90, end=-90 + 360 * positive / total, fill="#3a8f6b")
    text_center(draw, (cx, cy), f"{total} 个视频", font(36, True))

    draw.rectangle([760, 300, 792, 332], fill="#3a8f6b")
    draw.text((815, 292), f"检出手部关键点：{positive} 个（{positive / total:.1%}）", font=label_font, fill="#22313f")
    draw.rectangle([760, 365, 792, 397], fill="#b35f5f")
    draw.text((815, 357), f"未检出手部关键点：{zero} 个（{zero / total:.1%}）", font=label_font, fill="#22313f")
    draw.text((760, 455), f"平均 hand_ratio：{summary['avg_hand_ratio']:.4f}", font=small_font, fill="#52606d")
    draw.text((760, 500), f"平均 FPS：{summary['avg_fps']:.4f}", font=small_font, fill="#52606d")
    img.save(path)


def generate_images() -> list[tuple[str, Path, str]]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    perf = ASSET_DIR / "input_mode_performance.png"
    yolo = ASSET_DIR / "yolo_bridge_detection_distribution.png"
    reject = ASSET_DIR / "dataset_rejection_reasons.png"
    f1 = ASSET_DIR / "gesture_replay_f1.png"
    finger = ASSET_DIR / "finger_to_palm_ratio_examples.png"
    draw_performance(perf)
    draw_yolo(yolo)
    rejection = read_csv("dataset_rejection_reasons.csv")
    draw_bar_chart(
        reject,
        "Jester 样本筛选拒绝原因统计",
        [r["reason"] for r in rejection],
        [int(r["count"]) for r in rejection],
        "样本数",
        "#6d7f9d",
    )
    f1_rows = read_csv("gesture_precision_recall_f1.csv")
    draw_bar_chart(
        f1,
        "外部模板回放动态手势识别 F1",
        [r["gesture"].replace("motion_", "") for r in f1_rows],
        [float(r["f1"]) for r in f1_rows],
        "F1",
        "#4a7c59",
    )
    feature_rows = read_csv("finger_level_feature_summary.csv")
    wanted = ["ext_two_finger_spread_easy", "ext_motion_right_right_short", "ext_motion_up_right_short", "ext_horizontal_wave_easy"]
    picked = []
    seen = set()
    for row in feature_rows:
        if row["gesture_id"] in wanted and row["gesture_id"] not in seen:
            picked.append(row)
            seen.add(row["gesture_id"])
    draw_bar_chart(
        finger,
        "手指级变化与掌心轨迹比例示例",
        [r["gesture_id"].replace("ext_", "").replace("_easy", "") for r in picked],
        [float(r["finger_to_palm_path_ratio"]) for r in picked],
        "指尖路径/掌心路径",
        "#8a6fb0",
    )
    return [
        ("图 6-2 三种输入模式实时性能对比", perf, "图 6-2 表明，Mock 模式作为开发与答辩兜底链路具有最高且最稳定的帧率；ExternalBridge 模式在引入 UDP 外部视觉输入后仍保持较高帧率，并能记录包间隔与估算延迟；Native MediaPipe 模式受 Unity 内部视觉处理开销影响，平均帧率相对较低。"),
        ("图 6-3 YOLO+MediaPipe 外部桥接有效检出分布", yolo, "图 6-3 用于说明 YOLO+MediaPipe 链路的当前边界：外部程序可以完成视频处理流程，但有效手部关键点检出比例仍偏低，因此本文仅将其作为外部桥接可行性和后续优化方向。"),
        ("图 6-4 Jester 样本筛选拒绝原因统计", reject, "图 6-4 说明离线样本进入模板回放前需要经过规则筛选，主要拒绝原因包括无动作标签、未检出手部和运动分数不足，这也解释了实验结果不能简单等同于真实用户长期测试。"),
        ("图 6-5 外部模板回放动态手势识别 F1", f1, "图 6-5 显示，在严格筛选且模板一致的离线回放子集上，八类方向性动态手势均能被当前规则识别器稳定识别。"),
        ("图 6-6 手指级变化与掌心轨迹比例示例", finger, "图 6-6 表明，双指外滑、响指等动作更依赖指尖距离、峰值速度和振荡次数等手指级特征，不能完全依靠掌心位移判断。"),
    ]


def find_insert_index(doc: Document) -> int:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith("第 7 章") or paragraph.text.strip().startswith("第7章"):
            return index
    raise ValueError("Cannot find chapter 7 anchor")


def insert_paragraph_before(paragraph, text: str, style=None):
    new = paragraph.insert_paragraph_before(text)
    if style is not None:
        new.style = style
    return new


def add_run_paragraph(anchor, text: str, style_source=None, first_line: bool = True):
    p = insert_paragraph_before(anchor, "", style_source.style if style_source is not None else None)
    run = p.add_run(text)
    run.font.size = Pt(12)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(24)
    return p


def insert_into_docx(images: list[tuple[str, Path, str]]) -> None:
    doc = Document(SOURCE)
    idx = find_insert_index(doc)
    anchor = doc.paragraphs[idx]
    style_source = doc.paragraphs[max(0, idx - 1)]

    heading = insert_paragraph_before(anchor, "6.8 补充数据统计与结论图表", style_source.style)
    heading.paragraph_format.first_line_indent = None
    overview = {r["metric"]: r["value"] for r in read_csv("experiment_overview.csv")}
    perf = read_csv("input_mode_performance_summary.csv")
    yolo = {r["metric"]: r["value"] for r in read_csv("yolo_bridge_summary.csv")}
    add_run_paragraph(
        anchor,
        f"为进一步增强实验结论的可读性，本文根据实验补充数据对样本筛选、输入模式性能、YOLO 外部桥接和手指级动态特征进行统计可视化。补充实验从 Jester 样本中挖掘 {overview['mined_rows']} 条记录，接受 {overview['accepted_rows']} 条，接受率为 {overview['accepted_rate']}；AVI-200 扩展回放共采样 {overview['sampled_frames']} 帧，其中 MediaPipe 有效检出 {overview['detected_frames']} 帧，有效帧比例为 {overview['detected_frame_rate']}。",
        style_source,
    )
    add_run_paragraph(
        anchor,
        f"在性能方面，三种输入模式共统计 {overview['performance_runs']} 条运行记录。Mock 平均 FPS 为 {perf[0]['average_fps_mean']}，Native MediaPipe 平均 FPS 为 {perf[1]['average_fps_mean']}，ExternalBridge 平均 FPS 为 {perf[2]['average_fps_mean']}；ExternalBridge 平均包间隔为 {perf[2]['avg_packet_interval_ms_mean']} ms，估算链路延迟为 {perf[2]['avg_estimated_latency_ms_mean']} ms。",
        style_source,
    )

    for caption, path, note in images:
        p_img = insert_paragraph_before(anchor, "")
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(str(path), width=Inches(5.4))
        p_cap = insert_paragraph_before(anchor, caption, style_source.style)
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.first_line_indent = None
        add_run_paragraph(anchor, note, style_source)

    add_run_paragraph(
        anchor,
        f"YOLO 外部桥接实验中，18 个参考视频均完成处理，其中 {yolo['hand_positive_videos']} 个视频检出手部关键点，{yolo['zero_hand_videos']} 个视频未检出，平均 hand_ratio 为 {yolo['avg_hand_ratio']}，平均 FPS 为 {yolo['avg_fps']}。因此，YOLO+MediaPipe 当前应表述为外部桥接可行性和后续优化方向，不能表述为已完成高精度 YOLO 动态手势分类器。",
        style_source,
    )

    doc.save(TARGET)
    doc.save(ARCHIVE_COPY)


def main() -> None:
    images = generate_images()
    insert_into_docx(images)
    print(TARGET)
    print(ARCHIVE_COPY)


if __name__ == "__main__":
    main()
