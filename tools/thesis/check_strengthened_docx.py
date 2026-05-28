from pathlib import Path

from docx import Document


path = Path("毕业设计提交归档-曹逸天") / "3毕业设计说明书-曹逸天-内容风险补强版.docx"
doc = Document(path)
print("paragraphs", len(doc.paragraphs), "tables", len(doc.tables), "images", len(doc.inline_shapes))
for needle in ["表 2-1", "算法 4-1", "本章实验并不将目标设定"]:
    print("---", needle)
    for para in doc.paragraphs:
        if needle in para.text:
            print(para.text[:200])
