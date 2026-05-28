from pathlib import Path

from docx import Document


docx_path = Path("毕业设计提交归档-曹逸天") / "3毕业设计说明书-曹逸天.docx"
doc = Document(docx_path)

patterns = [
    "注意：",
    "抄袭",
    "本论文给出的内容仅供参考",
    "```",
    "public ",
    "private ",
    "class ",
    "void ",
]

for index, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    if any(pattern in text for pattern in patterns):
        print(index, repr(text))

print("paragraphs", len(doc.paragraphs), "tables", len(doc.tables), "inline_shapes", len(doc.inline_shapes))
